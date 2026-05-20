from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document

FIXTURES_DIR = Path(__file__).parent / "fixtures"


@pytest.fixture(scope="session")
def fixtures_dir() -> Path:
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    return FIXTURES_DIR


@pytest.fixture(scope="session")
def sample_pdf(fixtures_dir: Path) -> Path:
    path = fixtures_dir / "sample_paper.pdf"
    if path.exists():
        return path

    doc = fitz.open()
    page = doc.new_page()
    text = """Sample Paper Title

Abstract
This paper studies citation parsing for reference checks.

Keywords: parsing, references, doi

1. Introduction
Previous work has shown important results [1] and follow-up [2-3].
Smith et al. also reported similar findings (Jones, 2019).

2. Methods
We applied standard procedures [1].

References
[1] Smith J. Example title. Journal A. 2020;1(1):1-10. doi:10.1000/example.1
[2] Doe A. Another study. Science. 2019;5:100-110. doi:10.1000/example.2
"""
    page.insert_text((72, 72), text, fontsize=11)
    doc.save(path)
    doc.close()
    return path


@pytest.fixture(scope="session")
def sample_docx(fixtures_dir: Path) -> Path:
    path = fixtures_dir / "sample_paper.docx"
    if path.exists():
        return path

    document = Document()
    document.add_heading("Sample Paper Title", level=0)
    document.add_paragraph("Abstract")
    document.add_paragraph(
        "This paper studies citation parsing for reference checks."
    )
    document.add_paragraph("Keywords: parsing, references, doi")
    document.add_paragraph("1. Introduction")
    document.add_paragraph(
        "Previous work has shown important results [1] and follow-up [2]."
    )
    document.add_paragraph("References")
    document.add_paragraph(
        "[1] Smith J. Example title. Journal A. 2020. doi:10.1000/example.1"
    )
    document.save(path)
    return path
